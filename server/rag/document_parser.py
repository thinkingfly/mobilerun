"""文档解析模块 - 支持 Word 和 PDF 格式。"""

import logging
from pathlib import Path

logger = logging.getLogger("mobilerun.server")


def parse_word(file_path: str) -> str:
    """解析 Word 文档 (.docx)。"""
    from docx import Document

    doc = Document(file_path)
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text.strip())

    # 解析表格
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                paragraphs.append(row_text)

    return "\n\n".join(paragraphs)


def parse_pdf(file_path: str) -> str:
    """解析 PDF 文档。"""
    from PyPDF2 import PdfReader

    reader = PdfReader(file_path)
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text and text.strip():
            pages.append(text.strip())

    return "\n\n".join(pages)


def parse_txt(file_path: str) -> str:
    """解析纯文本文件。"""
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()


def parse_document(file_path: str, filetype: str = None) -> str:
    """统一文档解析入口。

    Args:
        file_path: 文件路径
        filetype: 文件类型 (word/pdf/txt)，如果不指定则从扩展名推断

    Returns:
        解析后的文本内容
    """
    path = Path(file_path)
    suffix = path.suffix.lower()

    if filetype is None:
        if suffix in [".docx", ".doc"]:
            filetype = "word"
        elif suffix == ".pdf":
            filetype = "pdf"
        elif suffix == ".txt":
            filetype = "txt"
        else:
            raise ValueError(f"不支持的文件格式：{suffix}")

    if filetype == "word":
        return parse_word(file_path)
    elif filetype == "pdf":
        return parse_pdf(file_path)
    elif filetype == "txt":
        return parse_txt(file_path)
    else:
        raise ValueError(f"不支持的文件类型：{filetype}")
